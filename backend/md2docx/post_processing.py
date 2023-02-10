from docx import Document


doc = Document('samples/text.docx')
print(doc.part.numbering_part.numbering_definitions._numbering.xml)


for para in doc.paragraphs:
    if len(para._element.xpath('./w:pPr/w:numPr')) > 0 and para.style.name != 'Bibliography':
        print(para.text)
        pass


# print([p.text for p in doc.paragraphs ])


class PostProcessing:
    def __init__(self, file: str):
        self.doc = Document(file)


def main(file: str):
    post_processing = PostProcessing(file)
